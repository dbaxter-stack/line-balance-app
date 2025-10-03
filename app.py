
import streamlit as st
import pandas as pd
from collections import defaultdict
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Line Balance Reports", layout="wide")
st.title("📊 Line Balance Reports (Courses × Lines) — v5.2")

uploaded = st.file_uploader("Upload Student Allocations CSV (must include 'Code' and AL1..)", type=["csv"])

# ---------------- Data helpers ----------------
def melt_long(df):
    alloc_cols = [c for c in df.columns if str(c).startswith("AL")]
    if "Code" not in df.columns:
        st.error("CSV must include a 'Code' column for unique student identifiers.")
        st.stop()
    long = df.melt(
        id_vars=[c for c in df.columns if c not in alloc_cols],
        value_vars=alloc_cols,
        var_name="Line",
        value_name="Class",
    )
    long = long.dropna(subset=["Class"]).copy()
    long["Course"] = long["Class"].astype(str).str[:5]
    return long

def counts_from_long(long_df):
    return long_df.groupby(["Line","Course"]).size().reset_index(name="StudentCount")

def build_offerings(long_df):
    counts = counts_from_long(long_df)
    wide = counts.pivot(index="Course", columns="Line", values="StudentCount")
    course_to_lines = {course: [ln for ln, ct in row.dropna().items() if ct > 0] for course, row in wide.iterrows()}
    return wide, course_to_lines

def compute_imbalance(wide):
    rows = []
    for course, row in wide.iterrows():
        vals = row.dropna()
        nz = vals[vals > 0]
        appears_in = int((vals > 0).sum())
        if appears_in < 2:
            continue
        if len(nz) >= 2:
            rng = float(nz.max() - nz.min())
            rows.append({"Course": course, "Range": rng})
    return pd.DataFrame(rows).sort_values(["Range","Course"], ascending=[False, True]).reset_index(drop=True)

def build_student_schedule(long_df):
    sched = defaultdict(dict)
    for _, r in long_df.iterrows():
        sched[r["Code"]][r["Line"]] = r["Course"]
    return sched

# ---------------- Section-aware helpers ----------------
def get_course_sections_on_line(long_df, course, line):
    mask = (long_df["Course"] == course) & (long_df["Line"] == line)
    return sorted(long_df.loc[mask, "Class"].dropna().astype(str).unique().tolist())

def pick_destination_section(long_df, course, line):
    sections = get_course_sections_on_line(long_df, course, line)
    if not sections:
        return None
    counts = (
        long_df[(long_df["Course"] == course) & (long_df["Line"] == line)]
        .groupby("Class").size().reindex(sections, fill_value=0)
    )
    return counts.sort_values(kind="mergesort").index[0]

# ---------------- Planner (multi-step, safeguarded) ----------------
def plan_student_chain(student, course, from_ln, to_ln, sched, offerings, depth=2):
    # Destination free for this student?
    if to_ln not in sched[student]:
        return [(course, from_ln, to_ln)]
    if depth == 0:
        return None
    # Blocked by another course of this student
    blocking_course = sched[student][to_ln]
    # Try direct relocation of blocking_course
    for alt_ln in offerings.get(blocking_course, []):
        if alt_ln == to_ln or alt_ln in sched[student]:
            continue
        return [(blocking_course, to_ln, alt_ln), (course, from_ln, to_ln)]
    # Try two-step chain
    if depth > 1:
        for alt_ln in offerings.get(blocking_course, []):
            if alt_ln == to_ln:
                continue
            if alt_ln not in sched[student]:
                continue
            c2 = sched[student][alt_ln]
            for alt2 in offerings.get(c2, []):
                if alt2 in sched[student] or alt2 == alt_ln:
                    continue
                return [(c2, alt_ln, alt2), (blocking_course, to_ln, alt_ln), (course, from_ln, to_ln)]
    return None

def apply_chain_section_aware(long_df, sched, student, chain):
    # Validate all steps first
    for c, src_ln, dst_ln in chain:
        if sched[student].get(src_ln) != c:
            return False
        if dst_ln in sched[student]:
            return False
        dest_class = pick_destination_section(long_df, c, dst_ln)
        if dest_class is None:
            return False
    # Apply steps
    for c, src_ln, dst_ln in chain:
        dest_class = pick_destination_section(long_df, c, dst_ln)
        mask = (long_df["Code"] == student) & (long_df["Course"] == c) & (long_df["Line"] == src_ln)
        long_df.loc[mask, "Line"] = dst_ln
        long_df.loc[mask, "Class"] = dest_class
        sched[student].pop(src_ln, None)
        sched[student][dst_ln] = c
    return True

def compute_multi_move_plan_constrained(long_df, max_rounds=100, max_moves_per_student=3):
    sched = build_student_schedule(long_df)
    wide, offerings = build_offerings(long_df)
    moves = []
    improved = True
    rounds = 0
    moved_sc = set()  # (student, course) moved already
    student_move_counts = defaultdict(int)

    while improved and rounds < max_rounds:
        improved = False
        rounds += 1
        counts = counts_from_long(long_df)
        wide = counts.pivot(index="Course", columns="Line", values="StudentCount")
        for course, row in wide.iterrows():
            offered = [ln for ln, ct in row.dropna().items() if ct > 0]
            if len(offered) < 2:
                continue
            curr = {ln: int(row[ln]) for ln in offered}
            total = sum(curr.values())
            n = len(offered)
            base = total // n
            remainder = total % n
            lines_sorted_asc = sorted(offered, key=lambda ln: curr[ln])
            target = {ln: base for ln in offered}
            for ln in lines_sorted_asc[:remainder]:
                target[ln] = base + 1
            surplus = [ln for ln in offered if curr[ln] > target[ln]]
            deficit = [ln for ln in offered if curr[ln] < target[ln]]
            if not surplus or not deficit:
                continue
            for to_ln in deficit:
                need = target[to_ln] - curr[to_ln]
                if need <= 0:
                    continue
                for from_ln in surplus:
                    give = curr[from_ln] - target[from_ln]
                    if give <= 0:
                        continue
                    candidates = long_df[(long_df["Course"] == course) & (long_df["Line"] == from_ln)]["Code"].tolist()
                    moved_local = False
                    for student in candidates:
                        if student_move_counts[student] >= max_moves_per_student:
                            continue
                        if (student, course) in moved_sc:
                            continue
                        chain = plan_student_chain(student, course, from_ln, to_ln, sched, offerings, depth=2)
                        if chain is None:
                            continue
                        proposed_courses = [c for (c, _, _) in chain]
                        if any((student, c) in moved_sc for c in proposed_courses):
                            continue
                        if student_move_counts[student] + len(chain) > max_moves_per_student:
                            continue
                        ok = apply_chain_section_aware(long_df, sched, student, chain)
                        if not ok:
                            continue
                        # Record each step of the chain
                        for c, src_ln, dst_ln in chain:
                            moves.append({
                                "StudentCode": student,
                                "Course": c,
                                "FromLine": src_ln,
                                "ToLine": dst_ln
                            })
                            moved_sc.add((student, c))
                            student_move_counts[student] += 1
                        improved = True
                        moved_local = True
                        break
                    if moved_local:
                        break
                if improved:
                    break
            if improved:
                break
    return pd.DataFrame(moves), long_df

# ---------------- Impact & ranges ----------------
def build_impact(long_before, long_after):
    before = long_before.groupby(["Course","Line"]).size().reset_index(name="Before")
    after = long_after.groupby(["Course","Line"]).size().reset_index(name="After")
    impact = pd.merge(before, after, on=["Course","Line"], how="outer").fillna(0)
    impact["Before"] = impact["Before"].astype(int)
    impact["After"] = impact["After"].astype(int)
    impact["Change"] = impact["After"] - impact["Before"]
    return impact.sort_values(["Course","Line"]).reset_index(drop=True)

def build_ranges_from_impact_ignore_zeros(impact_df):
    rows = []
    for course, grp in impact_df.groupby("Course"):
        nb = grp["Before"][grp["Before"] > 0]
        na = grp["After"][grp["After"] > 0]
        rng_b = int(nb.max() - nb.min()) if not nb.empty else 0
        rng_a = int(na.max() - na.min()) if not na.empty else 0
        rows.append({"Course": course, "RangeBefore": rng_b, "RangeAfter": rng_a, "Improvement": rng_b - rng_a})
    df = pd.DataFrame(rows)
    # Only keep courses with positive range either before or after
    df = df[(df["RangeBefore"] > 0) | (df["RangeAfter"] > 0)].reset_index(drop=True)
    return df

# ---------------- Word export ----------------
def docx_from_reports(moves_df, impact_df) -> bytes:
    doc = Document()
    doc.add_heading("Student Move Suggestions & Impact Summary", level=1)

    impact_sorted = impact_df.sort_values(["Course","Line"]).reset_index(drop=True)
    ranges_df = build_ranges_from_impact_ignore_zeros(impact_sorted)

    # Quick summary
    total_moves = len(moves_df)
    courses_improved = int((ranges_df["Improvement"] > 0).sum())
    avg_improvement = float(ranges_df.loc[ranges_df["Improvement"] > 0, "Improvement"].mean()) if courses_improved > 0 else 0.0
    p = doc.add_paragraph("Quick Summary"); p.runs[0].bold = True
    for item in [
        f"Total moves proposed: {total_moves}",
        f"Courses with improved balance: {courses_improved}",
        f"Average improvement in course range (ignoring 0s): {avg_improvement:.1f}",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Courses still unbalanced (RangeAfter > 3), using filtered ranges
    doc.add_heading("Courses Still Unbalanced (Range > 3 After Moves, ignoring 0s)", level=2)
    alert_df = ranges_df[ranges_df["RangeAfter"] > 3].sort_values("RangeAfter", ascending=False)
    if not alert_df.empty:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Course"; hdr[1].text = "Range After"; hdr[2].text = "Range Before"
        for _, r in alert_df.iterrows():
            row = table.add_row().cells
            row[0].text = str(r["Course"]); row[1].text = str(int(r["RangeAfter"])); row[2].text = str(int(r["RangeBefore"]))
    else:
        doc.add_paragraph("All courses balanced within a range of 3.")

    doc.add_paragraph()

    # Per-course range summary (only positive ranges, zeros ignored)
    doc.add_heading("Per-course Range Summary (Only courses with range > 0; zeros ignored)", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Course"
    hdr[1].text = "Range Before"
    hdr[2].text = "Range After"
    hdr[3].text = "Improvement"
    for _, r in ranges_df.sort_values(["Improvement","Course"], ascending=[False, True]).iterrows():
        row = table.add_row().cells
        row[0].text = str(r["Course"])
        row[1].text = str(int(r["RangeBefore"]))
        row[2].text = str(int(r["RangeAfter"]))
        row[3].text = str(int(r["Improvement"]))

    # Student moves grouped
    doc.add_heading("Student Moves (Grouped by StudentCode)", level=2)
    if not moves_df.empty:
        msort = moves_df.sort_values(["StudentCode","Course","FromLine","ToLine"]).reset_index(drop=True)
        cur = None
        for _, r in msort.iterrows():
            sc = str(r["StudentCode"]); c = str(r["Course"]); fr = str(r["FromLine"]); to = str(r["ToLine"])
            if sc != cur:
                doc.add_heading(sc, level=3); cur = sc
            doc.add_paragraph(f"{c}: {fr} \u2192 {to}", style="List Bullet")
    else:
        doc.add_paragraph("No moves proposed.")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ---------------- UI ----------------
if uploaded is not None:
    df = pd.read_csv(uploaded)
    long = melt_long(df)
    counts = counts_from_long(long)
    wide, _ = build_offerings(long)
    imbalance = compute_imbalance(wide)

    st.subheader("1) Imbalanced Courses")
    st.dataframe(imbalance, use_container_width=True)

    st.subheader("2) Move Suggestions")
    enable_multi = st.toggle("Enable multi-step per-student planner (with safeguards)", value=True)
    max_moves = st.number_input("Max moves per student", min_value=1, max_value=10, value=3, step=1)
    if enable_multi:
        moves, long_after = compute_multi_move_plan_constrained(long.copy(), max_rounds=200, max_moves_per_student=int(max_moves))
    else:
        moves = pd.DataFrame(columns=["StudentCode","Course","FromLine","ToLine"])
        long_after = long.copy()
    st.dataframe(moves, use_container_width=True)

    st.subheader("3) Before / After Impact")
    impact = build_impact(long, long_after)
    st.dataframe(impact, use_container_width=True)

    # Downloads
    def df_to_csv(df): return df.to_csv(index=False).encode("utf-8")
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.download_button("Imbalanced Courses (CSV)", data=df_to_csv(imbalance), file_name="imbalanced_courses.csv")
    with c2:
        st.download_button("Move Suggestions (CSV)", data=df_to_csv(moves), file_name="move_suggestions.csv")
    with c3:
        st.download_button("Before-After Impact (CSV)", data=df_to_csv(impact), file_name="before_after_impact.csv")
    with c4:
        docx_bytes = docx_from_reports(moves, impact)
        st.download_button("Word Report (.docx)", data=docx_bytes, file_name="Student_Move_Suggestions_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Upload a CSV to start.")
