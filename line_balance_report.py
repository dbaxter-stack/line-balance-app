
#!/usr/bin/env python3
import argparse, os
from collections import defaultdict, deque

def get_course_sections_on_line(long_df, course, line):
    """Return the list of distinct class codes (sections) for a course on a given line."""
    mask = (long_df["Course"] == course) & (long_df["Line"] == line)
    return sorted(long_df.loc[mask, "Class"].dropna().astype(str).unique().tolist())

def pick_destination_section(long_df, course, line):
    """Pick the least-filled class section (Class code) for the course on the target line."""
    sections = get_course_sections_on_line(long_df, course, line)
    if not sections:
        return None
    counts = (
        long_df[(long_df["Course"] == course) & (long_df["Line"] == line)]
        .groupby("Class")
        .size()
        .reindex(sections, fill_value=0)
    )
    dest_class = counts.sort_values(kind="mergesort").index[0]
    return dest_class

from io import BytesIO
import pandas as pd
from docx import Document

def counts_from_long(long_df):
    return long_df.groupby(['Line','Course']).size().reset_index(name='StudentCount')

def build_offerings(long_df):
    counts = counts_from_long(long_df)
    wide = counts.pivot(index='Course', columns='Line', values='StudentCount')
    course_to_lines = {course: [ln for ln, ct in row.dropna().items() if ct > 0] for course, row in wide.iterrows()}
    return wide, course_to_lines

def build_student_schedule(long_df):
    sched = defaultdict(dict)
    for _, r in long_df.iterrows():
        sched[r['Code']][r['Line']] = r['Course']
    return sched

def compute_imbalance(wide):
    rows = []
    for course, row in wide.iterrows():
        vals = row.dropna()
        nz = vals[vals>0]
        appears_in = int((vals>0).sum())
        if appears_in < 2:
            continue
        if len(nz) >= 2:
            rng = float(nz.max()-nz.min())
            rows.append({'Course': course, 'Range': rng})
    return pd.DataFrame(rows).sort_values(['Range','Course'], ascending=[False, True]).reset_index(drop=True)

def plan_student_chain(student, course, from_ln, to_ln, sched, offerings, depth=2):
    if to_ln not in sched[student]:
        return [(course, from_ln, to_ln)]
    if depth == 0:
        return None
    blocking_course = sched[student][to_ln]
    for alt_ln in offerings.get(blocking_course, []):
        if alt_ln == to_ln or alt_ln in sched[student]:
            continue
        return [(blocking_course, to_ln, alt_ln), (course, from_ln, to_ln)]
    if depth > 1:
        for alt_ln in offerings.get(blocking_course, []):
            if alt_ln == to_ln:
                continue
            if alt_ln not in sched[student]:
                continue
            c2 = sched[student][alt_ln]
            for alt2 in offerings.get(c2, []):
                if alt2 in sched[student] or alt2 == alt_ln:
                    continue
                return [(c2, alt_ln, alt2), (blocking_course, to_ln, alt_ln), (course, from_ln, to_ln)]
    return None

def compute_multi_move_plan_constrained(long_df, max_rounds=100, max_moves_per_student=3):
    sched = build_student_schedule(long_df)
    wide, offerings = build_offerings(long_df)
    moves = []
    improved = True
    rounds = 0
    moved_sc = set()
    student_move_counts = defaultdict(int)
    while improved and rounds < max_rounds:
        improved = False
        rounds += 1
        counts = counts_from_long(long_df)
        wide = counts.pivot(index='Course', columns='Line', values='StudentCount')
        for course, row in wide.iterrows():
            offered = [ln for ln, ct in row.dropna().items() if ct > 0]
            if len(offered) < 2:
                continue
            curr = {ln: int(row[ln]) for ln in offered}
            total = sum(curr.values())
            n = len(offered)
            base = total // n
            remainder = total % n
            lines_sorted_asc = sorted(offered, key=lambda ln: curr[ln])
            target = {ln: base for ln in offered}
            for ln in lines_sorted_asc[:remainder]:
                target[ln] = base + 1
            surplus = [ln for ln in offered if curr[ln] > target[ln]]
            deficit = [ln for ln in offered if curr[ln] < target[ln]]
            if not surplus or not deficit:
                continue
            for to_ln in deficit:
                need = target[to_ln] - curr[to_ln]
                if need <= 0:
                    continue
                for from_ln in surplus:
                    give = curr[from_ln] - target[from_ln]
                    if give <= 0:
                        continue
                    candidates = long_df[(long_df['Course']==course) & (long_df['Line']==from_ln)]['Code'].tolist()
                    moved_local = False
                    for student in candidates:
                        if student_move_counts[student] >= max_moves_per_student:
                            continue
                        if (student, course) in moved_sc:
                            continue
                        chain = plan_student_chain(student, course, from_ln, to_ln, sched, offerings, depth=2)
                        if chain is None:
                            continue
                        proposed_courses = [c for (c, _, _) in chain]
                        if any((student, c) in moved_sc for c in proposed_courses):
                            continue
                        if student_move_counts[student] + len(chain) > max_moves_per_student:
                            continue
                        valid = True
                        for c, src_ln, dst_ln in chain:
                            if sched[student].get(src_ln) != c or dst_ln in sched[student]:
                                valid = False
                                break
                        if not valid:
                            continue
                        for c, src_ln, dst_ln in chain:
                            sched[student].pop(src_ln, None)
                            sched[student][dst_ln] = c
                            mask = (long_df['Code']==student) & (long_df['Course']==c) & (long_df['Line']==src_ln)
                            # Assign to specific destination section to preserve separate classes
dest_class = pick_destination_section(long_df, c, dst_ln)
if dest_class is None:
    valid = False
    break
long_df.loc[mask, 'Line'] = dst_ln
long_df.loc[mask, 'Class'] = dest_class

                            moves.append({'StudentCode': student, 'Course': c, 'FromLine': src_ln, 'ToLine': dst_ln})
                            moved_sc.add((student, c))
                            student_move_counts[student] += 1
                        improved = True
                        moved_local = True
                        break
                    if moved_local:
                        break
                if improved:
                    break
            if improved:
                break
    return pd.DataFrame(moves), long_df

def build_impact(long_before, long_after):
    before = long_before.groupby(['Course','Line']).size().reset_index(name='Before')
    after = long_after.groupby(['Course','Line']).size().reset_index(name='After')
    impact = pd.merge(before, after, on=['Course','Line'], how='outer').fillna(0)
    impact['Before'] = impact['Before'].astype(int)
    impact['After'] = impact['After'].astype(int)
    impact['Change'] = impact['After'] - impact['Before']
    return impact.sort_values(['Course','Line']).reset_index(drop=True)

def build_ranges_from_impact(impact_df):
    rows = []
    for course, grp in impact_df.groupby('Course'):
        b = grp[grp['Before'] > 0].groupby('Course')['Before'].agg(['min','max'])
        a = grp[grp['After'] > 0].groupby('Course')['After'].agg(['min','max'])
        rng_b = int(b['max'].iloc[0] - b['min'].iloc[0]) if not b.empty else 0
        rng_a = int(a['max'].iloc[0] - a['min'].iloc[0]) if not a.empty else 0
        rows.append({'Course': course, 'RangeBefore': rng_b, 'RangeAfter': rng_a, 'Improvement': rng_b - rng_a})
    return pd.DataFrame(rows).sort_values('Improvement', ascending=False).reset_index(drop=True)

def write_docx(moves_df, impact_df, out_path):
    from docx import Document
    doc = Document()
    doc.add_heading("Student Move Suggestions & Impact Summary", level=1)
    impact_sorted = impact_df.sort_values(['Course','Line']).reset_index(drop=True)
    ranges_df = build_ranges_from_impact(impact_sorted)
    ranges_df_filtered = ranges_df[(ranges_df['RangeBefore'] > 0) | (ranges_df['RangeAfter'] > 0)].reset_index(drop=True)
    total_moves = len(moves_df)
    courses_improved = int((ranges_df['Improvement'] > 0).sum())
    avg_improvement = float(ranges_df.loc[ranges_df['Improvement'] > 0, 'Improvement'].mean()) if courses_improved > 0 else 0.0
    p = doc.add_paragraph("Quick Summary"); p.runs[0].bold = True
    for item in [f"Total moves proposed: {total_moves}", f"Courses with improved balance: {courses_improved}", f"Average improvement in course range: {avg_improvement:.1f}"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Courses Still Unbalanced (Range > 3 After Moves)", level=2)
    alert_df = ranges_df_filtered[ranges_df_filtered['RangeAfter'] > 3].sort_values('RangeAfter', ascending=False)
    if not alert_df.empty:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells; hdr[0].text="Course"; hdr[1].text="Range After"; hdr[2].text="Range Before"
        for _, r in alert_df.iterrows():
            row = table.add_row().cells
            row[0].text = str(r['Course']); row[1].text = str(int(r['RangeAfter'])); row[2].text = str(int(r['RangeBefore']))
    else:
        doc.add_paragraph("All courses balanced within a range of 3.")
    doc.add_heading("Student Moves (Grouped by StudentCode)", level=2)
    if not moves_df.empty:
        msort = moves_df.sort_values(['StudentCode','Course','FromLine','ToLine']).reset_index(drop=True)
        cur = None
        for _, r in msort.iterrows():
            stud = str(r['StudentCode']); c = str(r['Course']); fr = str(r['FromLine']); to = str(r['ToLine'])
            if stud != cur:
                doc.add_heading(stud, level=3); cur = stud
            doc.add_paragraph(f"{c}: {fr} \u2192 {to}", style="List Bullet")
    else:
        doc.add_paragraph("No moves proposed.")
    doc.save(out_path)

def main():
    ap = argparse.ArgumentParser(description="Line Balance Reports with Word export.")
    ap.add_argument("--input", "-i", required=True)
    ap.add_argument("--outdir", "-o", default="./out")
    ap.add_argument("--multi-move", action="store_true")
    ap.add_argument("--max-moves-per-student", type=int, default=3)
    args = ap.parse_args()

    df = pd.read_csv(args.input)
    alloc_cols = [c for c in df.columns if str(c).startswith('AL')]
    long0 = df.melt(id_vars=[c for c in df.columns if c not in alloc_cols], value_vars=alloc_cols, var_name='Line', value_name='Class')
    long0 = long0.dropna(subset=['Class']).copy()
    long0['Course'] = long0['Class'].astype(str).str[:5]

    wide0, _ = build_offerings(long0)
    imb = compute_imbalance(wide0)

    if args.multi_move:
        moves, long_after = compute_multi_move_plan_constrained(long0.copy(), max_rounds=200, max_moves_per_student=args.max_moves_per_student)
    else:
        moves = pd.DataFrame(columns=['StudentCode','Course','FromLine','ToLine'])
        long_after = long0.copy()

    impact = build_impact(long0, long_after)

    os.makedirs(args.outdir, exist_ok=True)
    imb.to_csv(os.path.join(args.outdir, "imbalanced_courses.csv"), index=False)
    moves.to_csv(os.path.join(args.outdir, "move_suggestions.csv"), index=False)
    impact.to_csv(os.path.join(args.outdir, "before_after_impact.csv"), index=False)

    # Excel
    try:
        with pd.ExcelWriter(os.path.join(args.outdir, "line_balance_reports.xlsx"), engine="openpyxl") as xlw:
            imb.to_excel(xlw, sheet_name="ImbalancedCourses", index=False)
            moves.to_excel(xlw, sheet_name="MoveSuggestions", index=False)
            impact.to_excel(xlw, sheet_name="BeforeAfterImpact", index=False)
    except Exception as e:
        print("Excel output skipped:", e)

    # Word
    try:
        write_docx(moves, impact, os.path.join(args.outdir, "Student_Move_Suggestions_Report.docx"))
    except Exception as e:
        print("Word output skipped:", e)

if __name__ == "__main__":
    main()
